from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


OUT_DIR = Path("docs/curriculos")
PDF_OUT = OUT_DIR / "IURY_RUAN_CV_ANDROID_MOBILE.pdf"
DOCX_OUT = OUT_DIR / "IURY_RUAN_CV_ANDROID_MOBILE.docx"


CONTACT = "Teresina, Brasil | iuryruansc@gmail.com | linkedin.com/in/iuryruansc | github.com/iuryruansc"

CONTENT = {
    "name": "Iury Ruan Sousa Castro",
    "headline": "Desenvolvedor Android | Kotlin | Aplicativos Mobile",
    "summary": (
        "Desenvolvedor Android com foco em aplicativos mobile nativos, utilizando Kotlin, "
        "arquitetura MVVM, consumo de APIs REST e boas práticas modernas de desenvolvimento "
        "Android. Tenho interesse em soluções mobile para meios de pagamento, varejo e "
        "operações digitais, com atenção a código limpo, usabilidade, estabilidade e evolução "
        "contínua do produto."
    ),
    "skills": [
        ("Android nativo", "Kotlin, Java, Android SDK, AndroidX, ciclo de vida de Activities e Fragments"),
        ("Arquitetura e UI", "MVVM, ViewModel, Data Binding, organização em camadas e componentes reutilizáveis"),
        ("Dados e integrações", "Retrofit, Gson, Room, consumo de APIs REST, persistência local e tratamento de estados"),
        ("Assíncrono e qualidade", "Coroutines, boas práticas de legibilidade, Git/GitHub e revisão de código"),
        ("Domínio de interesse", "Aplicações mobile para pagamentos, varejo, SmartPOS e fluxos transacionais"),
    ],
    "projects": [
        (
            "Pokemon Center - Aplicativo Android em Kotlin",
            [
                "Aplicativo mobile nativo que consome a PokeAPI para busca, navegação e visualização de detalhes de Pokémon.",
                "Implementação com Retrofit, Gson, Coroutines e AndroidX ViewModel para integração de dados e atualização da interface.",
                "Organização do projeto com foco em MVVM, separação de responsabilidades e experiência de uso fluida.",
                "GitHub: github.com/iuryruansc/pokemon-center",
            ],
        ),
        (
            "Estudos e prática em aplicativos Android",
            [
                "Criação de telas, fluxos de navegação e componentes mobile com foco em clareza visual e manutenção.",
                "Uso de persistência local, integração com serviços externos e controle de estados de carregamento, sucesso e erro.",
                "Evolução contínua em práticas modernas do ecossistema Android e interesse direcionado para soluções de pagamento.",
            ],
        ),
        (
            "cs_pdv - Sistema de ponto de venda e operações financeiras",
            [
                "Projeto backend/web, não Android, voltado a ponto de venda, controle de vendas, estoque, clientes, relatórios financeiros e rotinas operacionais.",
                "GitHub: github.com/iuryruansc/cs_pdv",
            ],
        ),
    ],
    "education": [
        "Tecnólogo em Análise e Desenvolvimento de Sistemas - FMU | 2022 - 2024",
        "Bacharel em Engenharia Civil - CEUPI | 2014 - 2019",
    ],
    "languages": ["Inglês: C2 Proficient - EF SET 82/100"],
}


def ensure_out_dir() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)


def build_pdf() -> None:
    ensure_out_dir()
    styles = getSampleStyleSheet()
    accent = colors.HexColor("#1f6f78")
    dark = colors.HexColor("#222222")

    title = ParagraphStyle(
        "Title",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=19,
        leading=22,
        textColor=dark,
        alignment=TA_LEFT,
        spaceAfter=2,
    )
    headline = ParagraphStyle(
        "Headline",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=13,
        textColor=accent,
        spaceAfter=4,
    )
    contact = ParagraphStyle(
        "Contact",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#444444"),
        spaceAfter=8,
    )
    section = ParagraphStyle(
        "Section",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11.5,
        leading=14,
        textColor=accent,
        spaceBefore=7,
        spaceAfter=4,
    )
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.2,
        leading=12.1,
        textColor=dark,
        spaceAfter=3,
    )
    bullet = ParagraphStyle(
        "Bullet",
        parent=body,
        leftIndent=10,
        firstLineIndent=-6,
        bulletIndent=0,
        spaceAfter=1.6,
    )
    skill_left = ParagraphStyle("SkillLeft", parent=body, fontName="Helvetica-Bold", fontSize=8.8, leading=11.2)
    skill_right = ParagraphStyle("SkillRight", parent=body, fontSize=8.8, leading=11.2)
    project_title = ParagraphStyle(
        "ProjectTitle",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=9.4,
        leading=12,
        spaceBefore=3,
        spaceAfter=1,
    )

    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        rightMargin=1.55 * cm,
        leftMargin=1.55 * cm,
        topMargin=1.25 * cm,
        bottomMargin=1.25 * cm,
    )

    story = [
        Paragraph(CONTENT["name"], title),
        Paragraph(CONTENT["headline"], headline),
        Paragraph(CONTACT, contact),
        Paragraph("Resumo profissional", section),
        Paragraph(CONTENT["summary"], body),
        Paragraph("Competências técnicas mobile", section),
    ]

    skill_rows = []
    for label, value in CONTENT["skills"]:
        skill_rows.append([Paragraph(label, skill_left), Paragraph(value, skill_right)])
    table = Table(skill_rows, colWidths=[4.0 * cm, 14.0 * cm], hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#e9f4f5")),
                ("TEXTCOLOR", (0, 0), (-1, -1), dark),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#d7e5e7")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#bdd4d7")),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.extend([table, Paragraph("Projetos e experiência prática", section)])

    for name, items in CONTENT["projects"]:
        story.append(Paragraph(name, project_title))
        for item in items:
            story.append(Paragraph(item, bullet, bulletText="-"))

    story.append(Paragraph("Formação acadêmica", section))
    for item in CONTENT["education"]:
        story.append(Paragraph(item, bullet, bulletText="-"))

    story.append(Paragraph("Idiomas", section))
    for item in CONTENT["languages"]:
        story.append(Paragraph(item, bullet, bulletText="-"))

    story.append(Spacer(1, 3))
    doc.build(story)


def add_docx_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 111, 120)


def add_docx_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    run.font.size = Pt(9.5)


def build_docx() -> None:
    ensure_out_dir()
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name.paragraph_format.space_after = Pt(0)
    run = name.add_run(CONTENT["name"])
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(34, 34, 34)

    head = document.add_paragraph()
    head.paragraph_format.space_after = Pt(2)
    run = head.add_run(CONTENT["headline"])
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(31, 111, 120)

    contact = document.add_paragraph()
    contact.paragraph_format.space_after = Pt(8)
    run = contact.add_run(CONTACT)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(68, 68, 68)

    add_docx_heading(document, "Resumo profissional")
    paragraph = document.add_paragraph(CONTENT["summary"])
    paragraph.paragraph_format.line_spacing = 1.05
    paragraph.paragraph_format.space_after = Pt(3)

    add_docx_heading(document, "Competências técnicas mobile")
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in CONTENT["skills"]:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        cells[0].paragraphs[0].runs[0].bold = True
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.8)

    add_docx_heading(document, "Projetos e experiência prática")
    for name, items in CONTENT["projects"]:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(name)
        run.bold = True
        run.font.size = Pt(9.8)
        for item in items:
            add_docx_bullet(document, item)

    add_docx_heading(document, "Formação acadêmica")
    for item in CONTENT["education"]:
        add_docx_bullet(document, item)

    add_docx_heading(document, "Idiomas")
    for item in CONTENT["languages"]:
        add_docx_bullet(document, item)

    document.save(DOCX_OUT)


if __name__ == "__main__":
    build_pdf()
    build_docx()
    print(PDF_OUT.resolve())
    print(DOCX_OUT.resolve())
